#!/usr/bin/env python3
"""
Word Skill — Direct programmatic access to Word documents.
Layer 1 (Skill/API) — No UI automation needed.

Usage:
  python3 word-skill.py read <filepath>
  python3 word-skill.py read-tables <filepath>
  python3 word-skill.py replace <filepath> <placeholder> <value> [--output <path>]
  python3 word-skill.py replace-batch <filepath> --replacements '<json>' [--output <path>]
  python3 word-skill.py fill-table <filepath> --table-index <n> --data '<json>' [--output <path>]
  python3 word-skill.py info <filepath>

All output is JSON.
"""
import sys, json, argparse, os

def main():
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest='command')

    # read: dump all paragraphs as text
    p_read = sub.add_parser('read')
    p_read.add_argument('filepath')

    # read-tables: dump all tables as JSON
    p_rt = sub.add_parser('read-tables')
    p_rt.add_argument('filepath')

    # replace: find and replace a placeholder
    p_rep = sub.add_parser('replace')
    p_rep.add_argument('filepath')
    p_rep.add_argument('placeholder')  # e.g. "<<CustomerName>>"
    p_rep.add_argument('value')
    p_rep.add_argument('--output', default=None)

    # replace-batch: multiple replacements at once
    p_rb = sub.add_parser('replace-batch')
    p_rb.add_argument('filepath')
    p_rb.add_argument('--replacements', required=True)  # JSON: {"<<X>>": "value", ...}
    p_rb.add_argument('--output', default=None)

    # fill-table: fill a table with rows of data
    p_ft = sub.add_parser('fill-table')
    p_ft.add_argument('filepath')
    p_ft.add_argument('--table-index', type=int, required=True)
    p_ft.add_argument('--data', required=True)  # JSON array of row arrays
    p_ft.add_argument('--output', default=None)

    # info: paragraph count, table count, placeholders found
    p_info = sub.add_parser('info')
    p_info.add_argument('filepath')

    args = parser.parse_args()
    if not args.command:
        parser.print_help()
        sys.exit(1)

    try:
        from docx import Document
    except ImportError:
        print(json.dumps({"error": "python-docx not installed. Run: pip3 install python-docx"}))
        sys.exit(1)

    import re

    filepath = os.path.expanduser(args.filepath)
    if not os.path.exists(filepath):
        print(json.dumps({"error": f"File not found: {filepath}"}))
        sys.exit(1)

    try:
        if args.command == 'info':
            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            # Find placeholders like <<Something>> or {{something}}
            placeholders = list(set(re.findall(r'<<[^>]+>>|{{[^}]+}}', full_text)))
            tables_info = []
            for i, table in enumerate(doc.tables):
                headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
                tables_info.append({"index": i, "rows": len(table.rows), "cols": len(table.columns), "headers": headers})
            print(json.dumps({
                "file": filepath,
                "paragraphs": len(doc.paragraphs),
                "tables": tables_info,
                "placeholders": placeholders
            }))

        elif args.command == 'read':
            doc = Document(filepath)
            paragraphs = [{"index": i, "text": p.text, "style": p.style.name} for i, p in enumerate(doc.paragraphs) if p.text.strip()]
            print(json.dumps({"paragraphs": paragraphs}))

        elif args.command == 'read-tables':
            doc = Document(filepath)
            tables = []
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append({"index": i, "rows": rows})
            print(json.dumps({"tables": tables}))

        elif args.command == 'replace':
            doc = Document(filepath)
            count = 0
            for para in doc.paragraphs:
                if args.placeholder in para.text:
                    for run in para.runs:
                        if args.placeholder in run.text:
                            run.text = run.text.replace(args.placeholder, args.value)
                            count += 1
            # Also replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if args.placeholder in para.text:
                                for run in para.runs:
                                    if args.placeholder in run.text:
                                        run.text = run.text.replace(args.placeholder, args.value)
                                        count += 1
            output = args.output or filepath
            doc.save(os.path.expanduser(output))
            print(json.dumps({"replaced": args.placeholder, "with": args.value, "count": count, "saved": output}))

        elif args.command == 'replace-batch':
            doc = Document(filepath)
            replacements = json.loads(args.replacements)
            counts = {}
            for key, value in replacements.items():
                counts[key] = 0
                for para in doc.paragraphs:
                    if key in para.text:
                        for run in para.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, str(value))
                                counts[key] += 1
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if key in para.text:
                                    for run in para.runs:
                                        if key in run.text:
                                            run.text = run.text.replace(key, str(value))
                                            counts[key] += 1
            output = args.output or filepath
            doc.save(os.path.expanduser(output))
            print(json.dumps({"replacements": counts, "saved": output}))

        elif args.command == 'fill-table':
            doc = Document(filepath)
            data = json.loads(args.data)
            table = doc.tables[args.table_index]
            for i, row_data in enumerate(data):
                row_idx = i + 1  # skip header
                if row_idx < len(table.rows):
                    for j, val in enumerate(row_data):
                        if j < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[j].text = str(val)
                else:
                    # Add new row
                    new_row = table.add_row()
                    for j, val in enumerate(row_data):
                        if j < len(new_row.cells):
                            new_row.cells[j].text = str(val)
            output = args.output or filepath
            doc.save(os.path.expanduser(output))
            print(json.dumps({"table_index": args.table_index, "rows_filled": len(data), "saved": output}))

    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()
